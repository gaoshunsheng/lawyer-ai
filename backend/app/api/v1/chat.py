import json
import uuid

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.graphs.consult_graph import run_consultation_stream
from app.core.database import get_db
from app.core.dependencies import get_current_user
from app.models import User
from app.schemas.chat import (
    ChatMessageResponse,
    ChatMessageSend,
    ChatSessionCreate,
    ChatSessionResponse,
    LinkCaseRequest,
)
from app.services import chat_service
from app.services.file_parser import parse_file

MAX_ATTACHMENT_SIZE = 10 * 1024 * 1024  # 10MB

router = APIRouter(prefix="/chat", tags=["chat"])


@router.post("/sessions", response_model=ChatSessionResponse, status_code=201)
async def create_chat_session(
    req: ChatSessionCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    session = await chat_service.create_session(
        db, current_user.id, current_user.tenant_id, req.title, req.case_id,
    )
    return ChatSessionResponse.model_validate(session)


@router.get("/sessions", response_model=list[ChatSessionResponse])
async def list_sessions(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    sessions = await chat_service.get_user_sessions(db, current_user.id)
    return [ChatSessionResponse.model_validate(s) for s in sessions]


@router.get("/sessions/{session_id}/messages", response_model=list[ChatMessageResponse])
async def get_messages(
    session_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    messages = await chat_service.get_session_messages(db, session_id)
    return [ChatMessageResponse.model_validate(m) for m in messages]


@router.post("/sessions/{session_id}/messages")
async def send_message(
    session_id: uuid.UUID,
    req: ChatMessageSend,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    history = await chat_service.get_session_messages(db, session_id)
    messages = [{"role": m.role, "content": m.content} for m in history]

    # Inject attachment text if provided
    user_content = req.content
    if req.attachment_text:
        user_content = f"【用户上传文件内容】\n{req.attachment_text}\n\n【用户问题】\n{req.content}"

    await chat_service.save_message(db, session_id, "user", user_content)

    # Build case context if linked
    case_context = ""
    case = await chat_service.get_linked_case(db, session_id)
    if case:
        case_context = f"\n\n关联案件信息：案件编号={case.case_number}, 类型={case.case_type}, 原告={case.plaintiff}, 被告={case.defendant}, 案件状态={case.status}"

    async def stream_response():
        full_answer = ""
        async for chunk in run_consultation_stream(
            user_content,
            tenant_id=current_user.tenant_id,
            session_id=session_id,
            user_id=current_user.id,
            messages=messages,
            case_context=case_context,
        ):
            full_answer += chunk
            yield f"data: {json.dumps({'content': chunk}, ensure_ascii=False)}\n\n"

        await chat_service.save_message(db, session_id, "assistant", full_answer)
        yield "data: [DONE]\n\n"

    return StreamingResponse(stream_response(), media_type="text/event-stream")


@router.post("/sessions/{session_id}/attachments")
async def upload_attachment(
    session_id: uuid.UUID,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    content = await file.read()
    if len(content) > MAX_ATTACHMENT_SIZE:
        raise HTTPException(413, "文件大小超过10MB限制")

    parsed = await parse_file(content, file.filename or "", file.content_type or "")
    if not parsed:
        raise HTTPException(400, "无法解析文件内容")

    attachment_meta = {
        "filename": file.filename,
        "content_type": file.content_type,
        "size": len(content),
        "parsed_text_length": len(parsed),
    }
    await chat_service.save_message(
        db, session_id, "system", f"用户上传了文件: {file.filename}",
        attachments=attachment_meta,
    )
    return {"parsed_text": parsed, "filename": file.filename}


@router.post("/sessions/{session_id}/export")
async def export_session(
    session_id: uuid.UUID,
    format: str = Query("docx", pattern="^(docx|txt)$"),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    messages = await chat_service.get_session_messages(db, session_id)
    if not messages:
        raise HTTPException(404, "会话没有消息")

    session = await chat_service.get_session(db, session_id)
    title = session.title if session else "咨询记录"

    if format == "docx":
        from io import BytesIO
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading(title, level=1)
        for msg in messages:
            prefix = {"user": "用户", "assistant": "AI助手", "system": "系统"}.get(msg.role, msg.role)
            doc.add_paragraph(f"[{prefix}] {msg.content}")
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=chat_{session_id}.docx"},
        )
    else:
        import urllib.parse
        lines = [f"=== {title} ===\n"]
        for msg in messages:
            prefix = {"user": "用户", "assistant": "AI助手", "system": "系统"}.get(msg.role, msg.role)
            lines.append(f"[{prefix}] {msg.content}\n")
        text = "\n".join(lines)
        safe_name = urllib.parse.quote(f"chat_{session_id}.txt")
        return StreamingResponse(
            iter([text.encode("utf-8")]),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
        )


@router.post("/sessions/{session_id}/link-case")
async def link_case(
    session_id: uuid.UUID,
    req: LinkCaseRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    session = await chat_service.link_case(db, session_id, req.case_id)
    if not session:
        raise HTTPException(404, "会话不存在")
    return {"message": "已关联案件", "case_id": str(req.case_id)}


@router.delete("/sessions/{session_id}/link-case")
async def unlink_case(
    session_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    session = await chat_service.unlink_case(db, session_id)
    if not session:
        raise HTTPException(404, "会话不存在")
    return {"message": "已取消关联案件"}
